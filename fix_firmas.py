"""
Reemplaza el párrafo de firmas por una tabla 2x2 con layout correcto:
Izquierda: título arriba / nombre abajo | Derecha: título arriba / nombre abajo
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm
from copy import deepcopy
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

doc = Document('PROTOCOLO_TEMPLATE.docx')

# ── Encontrar el párrafo de firmas ─────────────────────────────────────────────
firma_idx = None
for i, p in enumerate(doc.paragraphs):
    if 'firma_resp_contacto' in p.text:
        firma_idx = i
        firma_para = p
        break

if firma_idx is None:
    print("ERROR: No se encontró el párrafo de firmas")
    exit(1)

print(f"Párrafo de firmas encontrado en P{firma_idx}")

# ── Copiar el estilo de fuente del párrafo original ────────────────────────────
# Obtener el primer run para copiar color y fuente
original_runs = firma_para.runs
font_name = "Calibri"
font_size = Pt(10)
font_bold = True
if original_runs:
    r = original_runs[0]
    if r.font.name:
        font_name = r.font.name
    if r.font.size:
        font_size = r.font.size
    font_bold = r.font.bold

# ── Insertar una tabla de 1 fila x 2 columnas ANTES del párrafo de firmas ─────
# Acceder al body del XML
body = doc.element.body
firma_elem = firma_para._element

# Crear tabla 2 x 2 (fila título + fila nombre)
table = doc.add_table(rows=2, cols=2)
table.style = doc.styles['Normal Table'] if 'Normal Table' in [s.name for s in doc.styles] else None

# Eliminar bordes de la tabla para que quede limpia
def remove_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)

remove_borders(table)

# Función para aplicar estilo a un run
def style_run(run, bold=True, size=Pt(10)):
    run.font.name = font_name
    run.font.size = size
    run.font.bold = bold

# Fila 0: Títulos
row0 = table.rows[0]
row0.cells[0].paragraphs[0].clear()
r0_left = row0.cells[0].paragraphs[0].add_run("FIRMA DEL RESPONSABLE DEL CONTACTO")
style_run(r0_left, bold=True, size=Pt(9))

row0.cells[1].paragraphs[0].clear()
r0_right = row0.cells[1].paragraphs[0].add_run("FIRMA RESPONSABLE DE GESTIÓN")
style_run(r0_right, bold=True, size=Pt(9))

# Fila 1: Nombres (etiquetas Jinja)
row1 = table.rows[1]
row1.cells[0].paragraphs[0].clear()
r1_left = row1.cells[0].paragraphs[0].add_run("{{ firma_resp_contacto }}")
style_run(r1_left, bold=False, size=Pt(10))

row1.cells[1].paragraphs[0].clear()
r1_right = row1.cells[1].paragraphs[0].add_run("{{ firma_resp_gestion }}")
style_run(r1_right, bold=False, size=Pt(10))

# Añadir línea en blanco sobre la firma (espacio visual)
# Añadir línea en blanco en las celdas de nombres
def add_top_border_to_cell(cell):
    """Agrega un borde superior a la celda (simula línea de firma)"""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    tcBorders = OxmlElement('w:tcBorders')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '6')
    top.set(qn('w:space'), '0')
    top.set(qn('w:color'), '000000')
    tcBorders.append(top)
    tcPr.append(tcBorders)

add_top_border_to_cell(row1.cells[0])
add_top_border_to_cell(row1.cells[1])

# ── Mover la tabla al lugar del párrafo de firmas en el XML ───────────────────
tbl_elem = table._tbl
body.remove(tbl_elem)          # docx.add_table lo pone al final, lo removemos
body.insert(list(body).index(firma_elem), tbl_elem)  # lo insertamos en su lugar

# Eliminar el párrafo original de firmas
firma_elem.getparent().remove(firma_elem)

doc.save('PROTOCOLO_TEMPLATE.docx')
print("Tabla de firmas creada exitosamente.")
