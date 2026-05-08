"""
Segunda pasada: corrección quirúrgica de etiquetas en PROTOCOLO_TEMPLATE.docx.
Opera run-by-run para no alterar formato (negrita, color, fuente, tamaño).
Busca el texto completo del párrafo y reemplaza solo el texto que cambia dentro de los runs.
"""
from docx import Document
import re, sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

doc = Document('PROTOCOLO_TEMPLATE.docx')

# Mapa de reemplazos: (texto a buscar, texto a poner)
# Todos sobre texto plano del párrafo (p.text)
REPLACEMENTS = {
    # Corrección mayúsculas que Word dejó mal
    '{{ NOMBRE_CONTACTO }}':     '{{ nombre_contacto }}',
    '{{ RELACION_PACIENTE }}':   '{{ relacion_paciente }}',
    '{{ POLICLINICO }}':         '{{ policlinico }}',
    '{{ FIRMA_RESP_CONTACTO }}': '{{ firma_resp_contacto }}',
    '{{ FIRMA_RESP_GESTION }}':  '{{ firma_resp_gestion }}',

    # Campos que quedaron sin etiqueta (aún con líneas)
    'MUY BUENOS DÍAS (TARDES), MI NOMBRE ES (registrar responsable del Llamado)':
        'MUY BUENOS DÍAS (TARDES), MI NOMBRE ES: {{ responsable_llamado }}',
    
    # Fecha/hora llamados
    'FECHA Y HORA DEL LLAMADO 1        ___/____/_____             _______ hrs':
        'FECHA Y HORA DEL LLAMADO 1: {{ fecha_llamado1 }} {{ hora_llamada1 }} hrs',
    'FECHA Y HORA DEL LLAMADO 2       ___/____/_____             _______ hrs':
        'FECHA Y HORA DEL LLAMADO 2: {{ fecha_llamado2 }} {{ hora_llamado2 }} hrs',

    # Teléfonos
    'NUMEROS DE TELEFONO                 _________________         _________________        _________________':
        'NUMEROS DE TELEFONO: {{ telefono_paciente }} / {{ telefono_alternativo }}',

    # Nombre completo / RUT paciente
    'COMPLETO _____________________________________________________    RUT _______________________':
        'COMPLETO: {{ nombre_paciente }}    RUT: {{ rut_paciente }}',

    # Teléfono + Dirección en línea (sabe ubicar)
    'TELEFONO_____________________________   DIRECCIÓN __________________________________________':
        'TELEFONO: {{ telefono_paciente }}   DIRECCIÓN: {{ direccion }}',

    # Nombre completo receptor
    'COMPLETO  _________________________________________________________________________________':
        'COMPLETO: {{ nombre_receptor }}',

    # Diagnóstico (la línea que quedó incompleta)
    'DIAGNÓSTICO  _____________________________________________________________________________':
        'DIAGNÓSTICO: {{ diagnostico }}',

    # Observaciones
    'OBSERVACIONES: _______________________________________________________________________________________________________________________________________':
        'OBSERVACIONES: {{ observaciones }}',
}

def replace_in_paragraph(para, find, replace):
    """Reemplaza `find` en el texto completo del párrafo, distribuyendo el cambio
    sobre los runs para NO perder el formato."""
    full_text = ''.join(r.text for r in para.runs)
    if find not in full_text:
        return False
    new_text = full_text.replace(find, replace)
    # Poner todo en el primer run y vaciar los demás
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ''
    return True

# ── Párrafos normales ──────────────────────────────────────────────────────────
for p in doc.paragraphs:
    for find, repl in REPLACEMENTS.items():
        replace_in_paragraph(p, find, repl)

# ── Celdas de tablas ──────────────────────────────────────────────────────────
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for find, repl in REPLACEMENTS.items():
                    replace_in_paragraph(p, find, repl)

# ── También aplica a líneas dentro de prestador/fecha (doble aparición) ───────
# Estas líneas aparecen dos veces en el doc (causal 1 y causal 4)
MULTI_REPLACE = {
    '__________________________________________________________': '{{ prestador }}/{{ fecha_atencion }}',
}
# Solo reemplazamos en los párrafos que tengan 'Establecimiento' o 'Fecha otorgamiento' en el párrafo anterior
prev_text = ''
for p in doc.paragraphs:
    cur = ''.join(r.text for r in p.runs)
    if 'Establecimiento prestador' in cur and '__________' in cur:
        replace_in_paragraph(p, '\t__________________________________________________________', ': {{ prestador }}')
    if 'Fecha otorgamiento' in cur and '__________' in cur:
        replace_in_paragraph(p, '\t\t__________________________________________________________', ': {{ fecha_atencion }}')
    prev_text = cur

doc.save('PROTOCOLO_TEMPLATE.docx')
print("Segunda pasada OK. Template listo.")
