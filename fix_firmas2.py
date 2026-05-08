"""
Reconstruye la tabla de firmas del PROTOCOLO_TEMPLATE.docx:
- Fila 0: espacio en blanco (para que el firmante escriba a mano)
- Fila 1: línea superior individual por celda (borde top solo en fila 1)
- Fila 2: título de la sección
- Fila 3: etiqueta Jinja del nombre
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

doc = Document('PROTOCOLO_TEMPLATE.docx')

# ── Eliminar tabla de firmas existente ────────────────────────────────────────
body = doc.element.body
tables = doc.tables
if tables:
    # La tabla de firmas es la primera (y única) tabla del doc
    tbl = tables[0]._tbl
    tbl.getparent().remove(tbl)
    print("Tabla anterior eliminada.")

# ── Crear nueva tabla 3 filas x 2 columnas ───────────────────────────────────
table = doc.add_table(rows=3, cols=2)

def set_cell_border(cell, **kwargs):
    """
    Aplica bordes a una celda.
    kwargs: top, bottom, left, right, insideH, insideV
    valor: 'single' o 'none'
    """
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in kwargs.items():
        border = OxmlElement(f'w:{side}')
        if val == 'single':
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '6')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
        else:
            border.set(qn('w:val'), 'none')
            border.set(qn('w:sz'), '0')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
        tcBorders.append(border)
    # Remove existing tcBorders if any
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcBorders)

def set_row_height(row, height_cm):
    tr = row._tr
    trPr = tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = OxmlElement('w:trPr')
        tr.insert(0, trPr)
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))  # 1 cm = 567 twips
    trPr.append(trHeight)

def style_paragraph(para, text, bold=False, size_pt=10):
    para.clear()
    run = para.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size_pt)
    run.font.bold = bold

# ── Fila 0: espacio en blanco (para que el firmante escriba a mano) ───────────
row0 = table.rows[0]
set_row_height(row0, 1.5)
for ci in range(2):
    cell = row0.cells[ci]
    set_cell_border(cell, top='none', bottom='none', left='none', right='none',
                    insideH='none', insideV='none')
    style_paragraph(cell.paragraphs[0], '')

# ── Fila 1: línea superior por celda (la "línea de firma") ───────────────────
row1 = table.rows[1]
set_row_height(row1, 0.2)
for ci in range(2):
    cell = row1.cells[ci]
    set_cell_border(cell, top='single', bottom='none', left='none', right='none',
                    insideH='none', insideV='none')
    style_paragraph(cell.paragraphs[0], '')

# ── Fila 2: título de cada sección ───────────────────────────────────────────
row2 = table.rows[2]
for ci, titulo in enumerate(['FIRMA DEL RESPONSABLE DEL CONTACTO', 'FIRMA RESPONSABLE DE GESTIÓN']):
    cell = row2.cells[ci]
    set_cell_border(cell, top='none', bottom='none', left='none', right='none',
                    insideH='none', insideV='none')
    style_paragraph(cell.paragraphs[0], titulo, bold=True, size_pt=9)

# ── Agregar fila extra con el nombre (etiqueta Jinja) ────────────────────────
# Añadir fila manualmente al XML (docx no tiene add_row directamente)
row3_tr = OxmlElement('w:tr')
for ci, tag in enumerate(['{{ firma_resp_contacto }}', '{{ firma_resp_gestion }}']):
    tc = OxmlElement('w:tc')
    tcPr = OxmlElement('w:tcPr')
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'bottom', 'left', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        tcBorders.append(b)
    tcPr.append(tcBorders)
    tc.append(tcPr)
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '20')  # 10pt = 20 half-points
    rPr.append(sz)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = tag
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    p.append(r)
    tc.append(p)
    row3_tr.append(tc)
table._tbl.append(row3_tr)

# ── Mover tabla al final del documento (antes del último párrafo) ─────────────
tbl_elem = table._tbl
body.remove(tbl_elem)
body.append(tbl_elem)

doc.save('PROTOCOLO_TEMPLATE.docx')
print("Tabla de firmas reconstruida correctamente.")
